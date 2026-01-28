"""
Question Import Script
Parses past exam questions from .txt and .docx files and imports into database
"""
import os
import sys
import io
import re
from pathlib import Path
from typing import List, Dict, Optional
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, JSON, Boolean, Enum as SQLEnum
from sqlalchemy.orm import sessionmaker, declarative_base
from datetime import datetime
import enum

# Fix console encoding for Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

def safe_print(msg):
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode('ascii', 'replace').decode('ascii'))

# Database setup
DATABASE_URL = "sqlite:///./medical_app.db"
engine = create_engine(DATABASE_URL, echo=False)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()

class DifficultyLevel(str, enum.Enum):
    EASY = "easy"
    MEDIUM = "medium"
    HARD = "hard"

class QuestionType(str, enum.Enum):
    GENERATED = "generated"
    PAST_PAPER = "past_paper"

class Question(Base):
    __tablename__ = "questions"
    id = Column(Integer, primary_key=True, index=True)
    slide_id = Column(Integer, nullable=True)
    department = Column(String(255), nullable=True, index=True)
    topic = Column(String(500), nullable=True)
    question_text = Column(Text, nullable=False)
    correct_answer = Column(String(500), nullable=False)
    distractors = Column(JSON, nullable=False)
    explanation = Column(Text, nullable=True)
    difficulty = Column(SQLEnum(DifficultyLevel), default=DifficultyLevel.MEDIUM)
    question_type = Column(SQLEnum(QuestionType), default=QuestionType.PAST_PAPER)
    is_past_paper = Column(Boolean, default=True)
    source_file = Column(String(500), nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)

# Try to import docx
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. Install with: pip install python-docx")

def parse_docx_questions(file_path: str) -> List[Dict]:
    """Extract text from docx and parse questions"""
    if not HAS_DOCX:
        return []
    
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        content = "\n".join(full_text)
        return parse_content_questions(content, Path(file_path).name)
    except Exception as e:
        safe_print(f"  Error parsing docx {file_path}: {e}")
        return []

def parse_content_questions(content: str, source_name: str) -> List[Dict]:
    """Parse questions from string content"""
    questions = []
    content = content.replace('\r\n', '\n').replace('\t', ' ')
    
    # Split content into question blocks by "ANSWER:" or similar marker
    # Some files use "CEVAP:" or "ANS:"
    units = re.split(r'(?i)ANSWER:|CEVAP:', content)
    
    for i in range(len(units) - 1):
        unit = units[i]
        answer_part = units[i+1]
        
        # Letter match at start of answer_part
        ans_match = re.search(r'^\s*([A-E])', answer_part, re.IGNORECASE)
        if not ans_match: continue
        correct_letter = ans_match.group(1).upper()
        
        options = []
        last_pos = len(unit)
        for letter in ['E', 'D', 'C', 'B', 'A']:
            # Search for option marker: "a)", "a.", "a-"
            opt_pattern = fr'(?:\n|\s){letter}[\)\.\-]\s+'
            matches = list(re.finditer(opt_pattern, unit, re.IGNORECASE))
            if not matches:
                # Fallback for start of file
                opt_pattern = fr'^{letter}[\)\.\-]\s+'
                matches = list(re.finditer(opt_pattern, unit, re.IGNORECASE))
                
            if matches:
                m = matches[-1]
                opt_text = unit[m.end():last_pos].strip()
                options.append(opt_text)
                last_pos = m.start()
        
        if len(options) < 2: continue
        
        q_text = unit[:last_pos].strip()
        # Remove leading numbers like "1.", "2-", etc.
        q_text = re.sub(r'^\d+[\.\s\-)]+', '', q_text).strip()
        
        options.reverse() # back to A, B, C...
        letter_idx = ord(correct_letter) - ord('A')
        
        if letter_idx < len(options):
            correct_answer = options[letter_idx]
            distractors = [opt for j, opt in enumerate(options) if j != letter_idx]
            
            questions.append({
                "question_text": q_text,
                "correct_answer": correct_answer,
                "distractors": distractors,
                "source_file": source_name
            })
    return questions

def parse_txt_questions(file_path: str) -> List[Dict]:
    """Parse questions from a text file"""
    content = ""
    for enc in ['utf-8', 'cp1254', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue
    if not content: return []
    return parse_content_questions(content, Path(file_path).name)

def main():
    source_path = r"C:\Users\Lenovo\OneDrive\Desktop\tıp5. sınıf"
    if not Path(source_path).exists():
        safe_print(f"[ERROR] Kaynak klasor bulunamadi: {source_path}")
        return
    
    safe_print("[INFO] Veritabani baglantisi kuruluyor...")
    Base.metadata.create_all(bind=engine)
    db = SessionLocal()
    
    total_imported = 0
    try:
        all_files = list(Path(source_path).rglob("*.txt")) + list(Path(source_path).rglob("*.docx"))
        safe_print(f"[INFO] Toplam {len(all_files)} dosya bulundu.")
        
        for file_path in all_files:
            name_lower = file_path.name.lower()
            parent_lower = file_path.parent.name.lower()
            
            if any(x in name_lower or x in parent_lower for x in ["soru", "sınav", "çıkmış", "quiz", "taramalar", "km"]):
                parts = file_path.parts
                dept_name = "Bilinmiyor"
                for i, part in enumerate(parts):
                    if "tıp5. sınıf" in part.lower() and i + 1 < len(parts):
                        dept_name = parts[i+1]
                        break
                
                safe_print(f"  Analiz ediliyor: {file_path.name} (Departman: {dept_name})")
                
                if file_path.suffix.lower() == '.txt':
                    parsed = parse_txt_questions(str(file_path))
                else:
                    parsed = parse_docx_questions(str(file_path))
                
                if not parsed:
                    safe_print(f"    [!] Soru bulunamadi or format uyumsuz.")
                    continue
                
                safe_print(f"    [OK] {len(parsed)} soru bulundu.")
                for q_data in parsed:
                    # Check if question already exists to avoid duplicates
                    # Simple check by text
                    exists = db.query(Question).filter(Question.question_text == q_data["question_text"]).first()
                    if exists: continue
                    
                    q = Question(
                        department=dept_name,
                        question_text=q_data["question_text"],
                        correct_answer=q_data["correct_answer"],
                        distractors=q_data["distractors"],
                        source_file=str(file_path.relative_to(source_path))
                    )
                    db.add(q)
                    total_imported += 1
                db.commit()
        safe_print(f"\n[SUCCESS] Toplam: {total_imported} yeni soru aktarildi!")
    except Exception as e:
        safe_print(f"[ERROR] Hata: {e}")
        db.rollback()
    finally:
        db.close()

if __name__ == "__main__":
    main()
